# test/create_test_docs.py
from docx import Document
from pathlib import Path

def create_docs():
    base = Path("tests/test_data")
    base.mkdir(parents=True, exist_ok=True)

    # 文件1：包含关键词
    doc1 = Document()
    doc1.add_paragraph("This is a test document with keyword Alice inside.")
    doc1.save(base / "test1.docx")

    # 文件2：不包含关键词
    doc2 = Document()
    doc2.add_paragraph("This document has no keyword.")
    doc2.save(base / "test2.docx")

    # 空文件
    doc3 = Document()
    doc3.save(base / "empty.docx")

if __name__ == "__main__":
    create_docs()
